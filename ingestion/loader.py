"""
Document loader for various file formats.

Supports loading documents from local filesystem without
requiring external APIs or services.
"""

from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class DocumentLoader:
    """
    Base class for loading documents from various formats.
    
    This class provides the interface for loading documents
    in an offline, local manner.
    """
    
    def __init__(self, file_path: Union[str, Path]) -> None:
        """
        Initialize the document loader.
        
        Args:
            file_path: Path to the document file
        """
        self.file_path: Path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")
    
    def load(self) -> Dict[str, Any]:
        """
        Load document content from file.
        
        Returns:
            Dictionary containing document metadata and content
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file format is not supported
        """
        return {
            "file_path": str(self.file_path),
            "file_name": self.file_path.name,
            "file_size": self.file_path.stat().st_size,
            "text": self.extract_text()
        }
    
    def extract_text(self) -> str:
        """
        Extract plain text from the document.
        
        Returns:
            Extracted text content
        """
        raise NotImplementedError("Subclasses must implement extract_text()")


class PDFLoader(DocumentLoader):
    """Loader for PDF documents."""
    
    def extract_text(self) -> str:
        """
        Extract text from PDF document.
        
        Returns:
            Extracted text content
            
        Raises:
            ImportError: If PyPDF2 is not installed
            Exception: If PDF cannot be read
        """
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF loading. Install with: pip install PyPDF2")
        
        text_parts: List[str] = []
        
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise ValueError(f"Failed to read PDF file {self.file_path}: {e}")
        
        return "\n".join(text_parts)


class DOCXLoader(DocumentLoader):
    """Loader for DOCX documents."""
    
    def extract_text(self) -> str:
        """
        Extract text from DOCX document.
        
        Returns:
            Extracted text content
            
        Raises:
            ImportError: If python-docx is not installed
            Exception: If DOCX cannot be read
        """
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX loading. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(self.file_path)
            text_parts: List[str] = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file {self.file_path}: {e}")


class TextLoader(DocumentLoader):
    """Loader for plain text documents."""
    
    def extract_text(self) -> str:
        """
        Extract text from plain text document.
        
        Returns:
            File content as string
            
        Raises:
            Exception: If file cannot be read
        """
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(self.file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise ValueError(f"Failed to read text file {self.file_path}: {e}")
        except Exception as e:
            raise ValueError(f"Failed to read text file {self.file_path}: {e}")


def get_loader(file_path: Union[str, Path]) -> DocumentLoader:
    """
    Factory function to get appropriate loader based on file extension.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Appropriate DocumentLoader instance
        
    Raises:
        ValueError: If file format is not supported
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == '.pdf':
        return PDFLoader(file_path)
    elif suffix in ['.docx', '.doc']:
        return DOCXLoader(file_path)
    elif suffix in ['.txt', '.md', '.text']:
        return TextLoader(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Supported formats: .pdf, .docx, .txt, .md")

